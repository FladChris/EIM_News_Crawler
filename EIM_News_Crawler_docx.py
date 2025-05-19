import os
import requests
from datetime import date, datetime
from pathlib import Path
from bs4 import BeautifulSoup
from htmldocx import HtmlToDocx
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml.shared import qn, OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE


BASE_URL = 'https://www.eim.uni-paderborn.de'
NEWS_LIST_PATH = '/eim-news-list/seite-{}'


def add_hyperlink(paragraph, text: str, url: str):
    """
    Fügt einen Hyperlink in einem Word-Dokument hinzu.
    """
    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    run = paragraph.add_run()
    run._r.append(hyperlink)
    run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    run.font.underline = True
    return hyperlink


def ensure_directory(year: str):
    Path(year).mkdir(exist_ok=True)


def save_docx(year: str,
              article_date: date,
              filename: str,
              headline: str,
              content_blocks: list,
              article_url: str,
              image_urls: list,
              contact_cards: dict):
    doc = Document()
    doc.add_heading(headline, level=2)

    converter = HtmlToDocx()
    for block in content_blocks:
        clean_html = str(block).replace('\xad', '').replace('|', '')
        converter.add_html_to_document(clean_html, doc)

    # URL
    p = doc.add_paragraph('URL des Artikels: ')
    add_hyperlink(p, article_url, article_url)

    # Bilder
    if image_urls:
        for url in image_urls:
            p = doc.add_paragraph('URL des Bildes: ')
            add_hyperlink(p, url, url)
    else:
        doc.add_paragraph('Kein Bild vorhanden')

    # Kontakte
    if contact_cards:
        for idx, card in contact_cards.items():
            doc.add_paragraph(f'Kontakt {idx}:')
            p = doc.add_paragraph('Bild-URL: ')
            add_hyperlink(p, card['image_url'], card['image_url'])
            doc.add_paragraph(f"Name des Kontakts: {card['link_text']}")
            p = doc.add_paragraph('PM-URL: ')
            add_hyperlink(p, card['link_url'], card['link_url'])
    else:
        doc.add_paragraph('Kein Kontakt vorhanden')

    out_path = Path(year) / f"{article_date.isoformat()}_{filename}.docx"
    doc.save(out_path)


def fetch_soup(session: requests.Session, url: str):
    resp = session.get(url)
    resp.raise_for_status()
    return BeautifulSoup(resp.content, 'html.parser')


def parse_contacts(soup: BeautifulSoup) -> dict:
    cards = {}
    sections = soup.select('div.business-card.teaser.last')
    for idx, sec in enumerate(sections):
        img = sec.select_one('img')
        link = sec.select_one('a')
        if img and link:
            cards[idx] = {
                'image_url': BASE_URL + img['src'],
                'link_text': link.text.strip(),
                'link_url': BASE_URL + link['href']
            }
    return cards


def parse_images(soup: BeautifulSoup) -> list:
    urls = []
    for a in soup.select('div.mediaelement-image a'):
        href = a.get('href')
        if href:
            urls.append(f"{BASE_URL}/{href.strip()}" )
    return urls


def main():
    year_str = input('Aus welchem Jahr sollen die Berichte sein?\n').strip()
    try:
        year = int(year_str)
    except ValueError:
        print('Ungültiges Jahr!')
        return

    start, end = date(year, 1, 1), date(year, 12, 31)
    ensure_directory(year_str)

    session = requests.Session()
    page = 1
    count = 0

    while True:
        list_url = BASE_URL + NEWS_LIST_PATH.format(page)
        soup = fetch_soup(session, list_url)
        items = soup.select('div.news-list-item_body')
        if not items:
            break

        for item in items:
            span = item.select_one('span.news-list-item_date')
            if not span:
                continue
            art_date = datetime.strptime(span.text.strip(), '%d.%m.%Y').date()
            if not (start <= art_date <= end):
                continue

            headline_tag = item.select_one('h3.news-list-item_headline')
            link_tag = item.select_one('a')
            if not headline_tag or not link_tag:
                continue

            headline = headline_tag.text.strip()
            href = link_tag['href']
            slug = href.strip('/').replace('/', '_').replace('-single', '')
            article_url = BASE_URL + href

            detail_soup = fetch_soup(session, article_url)
            content = detail_soup.select('div.news-detail_content')

            image_urls = parse_images(detail_soup)
            contact_cards = parse_contacts(detail_soup)

            save_docx(year_str, art_date, slug, headline, content, article_url, image_urls, contact_cards)
            print(f"Gespeichert: {art_date} - {headline}")
            count += 1

        page += 1

    print(f'Anzahl der News: {count}')


if __name__ == '__main__':
    main()
